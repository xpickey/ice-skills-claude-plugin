#!/usr/bin/env python3
"""
build_docx.py — generic DOCX builder for sales agents.

Usage:
    python3 build_docx.py spec.json out.docx

spec.json schema:
{
  "title": "...",
  "subtitle": "...",
  "blocks": [
    {"type": "h1", "text": "..."},
    {"type": "h2", "text": "..."},
    {"type": "h3", "text": "..."},
    {"type": "p", "text": "..."},
    {"type": "bullets", "items": ["...", "..."]},
    {"type": "numbered", "items": ["...", "..."]},
    {"type": "table", "headers": ["A","B"], "rows": [["x","y"]]},
    {"type": "page_break"},
    {"type": "callout", "text": "..."},
    {"type": "bilingual", "en": "...", "th": "..."}
  ]
}
"""
import sys, os, json, subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ⭐ นโยบายฟอนต์มาจาก SSOT เดียว — ห้าม hard-code ชื่อฟอนต์ในไฟล์นี้ (V02R01)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from font_policy import RAILS, resolve_font_policy, infer_rail   # noqa: E402


# ── ⭐ W1-W2 FONT BINDING ที่ระดับราก (V02R01 — 2026.08.04) ──────────────────
# บั๊กที่แก้: ไฟล์นี้ไม่มี font key ใน spec เลย และไม่เคยตั้ง w:rFonts
# → run ไทยไม่มี w:cs → Word ถอยไป Times New Roman (ไม่มี glyph ไทย)
# → substitute เงียบเป็น Angsana/Cordia = ตระกูลที่ทำลาย สระอำ 100%
# ตั้งที่ docDefaults + ทุก style ที่ใช้จริง → run ที่ไม่มี direct formatting inherit เอง
def _set_rfonts(rpr, font):
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(slot), font)          # W1 — ทั้ง 4 slot ตัวเดียวกัน (สเปก MS ขัดกันเองเรื่อง Thai)
    # ลบ theme attribute ที่จะ override ค่าเราได้
    for th in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rf.get(qn(th)) is not None:
            del rf.attrib[qn(th)]


def apply_font(doc, font, size_pt):
    # docDefaults — รากของทุก style · python-docx ไม่มี helper ให้ ต้องประกอบเอง
    # ⚠ ลำดับสำคัญ: w:docDefaults ต้องเป็น **ลูกตัวแรก** ของ w:styles ตาม schema
    #   (บทเรียนเดียวกับ CT_Settings ใน settings.xml — วางผิดที่ = Word สั่ง Repair ทั้งไฟล์)
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults"); styles_el.insert(0, dd)
    rprd = dd.find(qn("w:rPrDefault"))
    if rprd is None:
        rprd = OxmlElement("w:rPrDefault"); dd.insert(0, rprd)
    rpr = rprd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr"); rprd.append(rpr)
    _set_rfonts(rpr, font)
    for tag, val in (("w:sz", str(int(size_pt * 2))), ("w:szCs", str(int(size_pt * 2)))):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); rpr.append(el)
        el.set(qn("w:val"), val)        # w:szCs = ขนาดของ complex script (ไทย) — แยกจาก w:sz
    n = 1
    for st in doc.styles:
        try:
            _set_rfonts(st.element.get_or_add_rPr(), font); n += 1
        except Exception:
            pass
    # W2 — w:b/w:i ไม่มีผลกับ complex script ต้องมี bCs/iCs คู่เสมอ
    for p in doc.paragraphs:
        for r in p.runs:
            rp = r._r.get_or_add_rPr()
            for a, b in (("w:b", "w:bCs"), ("w:i", "w:iCs")):
                if rp.find(qn(a)) is not None and rp.find(qn(b)) is None:
                    rp.append(OxmlElement(b))
    return n


def add_callout(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_bilingual(doc, en, th):
    p = doc.add_paragraph()
    r1 = p.add_run("EN: ")
    r1.bold = True
    p.add_run(en)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("TH: ")
    r2.bold = True
    p2.add_run(th)


def build(spec_path, out_path):
    with open(spec_path) as f:
        spec = json.load(f)
    doc = Document()
    if spec.get("title"):
        h = doc.add_heading(spec["title"], 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if spec.get("subtitle"):
        p = doc.add_paragraph(spec["subtitle"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for block in spec.get("blocks", []):
        t = block.get("type")
        if t == "h1":
            doc.add_heading(block["text"], 1)
        elif t == "h2":
            doc.add_heading(block["text"], 2)
        elif t == "h3":
            doc.add_heading(block["text"], 3)
        elif t == "p":
            doc.add_paragraph(block["text"])
        elif t == "bullets":
            for item in block.get("items", []):
                doc.add_paragraph(item, style="List Bullet")
        elif t == "numbered":
            for item in block.get("items", []):
                doc.add_paragraph(item, style="List Number")
        elif t == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            tbl.style = "Light Grid"
            for j, h in enumerate(headers):
                tbl.rows[0].cells[j].text = str(h)
            for i, row in enumerate(rows, start=1):
                for j, v in enumerate(row):
                    tbl.rows[i].cells[j].text = str(v)
        elif t == "page_break":
            doc.add_page_break()
        elif t == "callout":
            add_callout(doc, block["text"])
        elif t == "bilingual":
            add_bilingual(doc, block.get("en", ""), block.get("th", ""))

    # ⭐ ฟอนต์มาจากราง (§3.0) · spec override ได้เมื่อลูกค้า/TOR บังคับ
    rail, _why = infer_rail(spec, out_path)
    print(f"📄 ราง: {rail}  ({_why})")
    if rail not in RAILS:
        sys.exit(f"rail ต้องเป็น {'|'.join(RAILS)} (ได้: {rail})")
    font = spec.get("font") or RAILS[rail]["font"]
    # ⭐ V02R02: ด่านนโยบายก่อนสร้างไฟล์ — ผิดนโยบาย = แก้ให้เป็นฟอนต์ราง + แจ้ง (ไม่ fail)
    font, _notices, _ = resolve_font_policy(font, rail, spec)
    for _n in _notices:
        print(_n)
    size = spec.get("font_size") or RAILS[rail]["size"]
    n = apply_font(doc, font, size)
    doc.save(out_path)
    print(f"OK: wrote {out_path} · font='{font}' {size}pt (rail={rail}) ผูก {n} style")

    subprocess.run([sys.executable, os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                                 "audit_fonts.py"), "--rail", rail,
                    *(["--allow-font", font] if spec.get("font") else []), out_path])


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: build_docx.py spec.json out.docx", file=sys.stderr)
        sys.exit(2)
    build(sys.argv[1], sys.argv[2])
