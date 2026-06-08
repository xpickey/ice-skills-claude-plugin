#!/usr/bin/env python3
"""
build_docx.py — generic DOCX builder for sales agents.

Usage:
    python3 build_docx.py spec.json out.docx

spec.json schema:
{
  "title": "...",
  "subtitle": "...",
  "blocks": [
    {"type": "h1", "text": "..."},
    {"type": "h2", "text": "..."},
    {"type": "h3", "text": "..."},
    {"type": "p", "text": "..."},
    {"type": "bullets", "items": ["...", "..."]},
    {"type": "numbered", "items": ["...", "..."]},
    {"type": "table", "headers": ["A","B"], "rows": [["x","y"]]},
    {"type": "page_break"},
    {"type": "callout", "text": "..."},
    {"type": "bilingual", "en": "...", "th": "..."}
  ]
}
"""
import sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_callout(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_bilingual(doc, en, th):
    p = doc.add_paragraph()
    r1 = p.add_run("EN: ")
    r1.bold = True
    p.add_run(en)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("TH: ")
    r2.bold = True
    p2.add_run(th)


def build(spec_path, out_path):
    with open(spec_path) as f:
        spec = json.load(f)
    doc = Document()
    if spec.get("title"):
        h = doc.add_heading(spec["title"], 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if spec.get("subtitle"):
        p = doc.add_paragraph(spec["subtitle"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for block in spec.get("blocks", []):
        t = block.get("type")
        if t == "h1":
            doc.add_heading(block["text"], 1)
        elif t == "h2":
            doc.add_heading(block["text"], 2)
        elif t == "h3":
            doc.add_heading(block["text"], 3)
        elif t == "p":
            doc.add_paragraph(block["text"])
        elif t == "bullets":
            for item in block.get("items", []):
                doc.add_paragraph(item, style="List Bullet")
        elif t == "numbered":
            for item in block.get("items", []):
                doc.add_paragraph(item, style="List Number")
        elif t == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            tbl.style = "Light Grid"
            for j, h in enumerate(headers):
                tbl.rows[0].cells[j].text = str(h)
            for i, row in enumerate(rows, start=1):
                for j, v in enumerate(row):
                    tbl.rows[i].cells[j].text = str(v)
        elif t == "page_break":
            doc.add_page_break()
        elif t == "callout":
            add_callout(doc, block["text"])
        elif t == "bilingual":
            add_bilingual(doc, block.get("en", ""), block.get("th", ""))

    doc.save(out_path)
    print(f"OK: wrote {out_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: build_docx.py spec.json out.docx", file=sys.stderr)
        sys.exit(2)
    build(sys.argv[1], sys.argv[2])
